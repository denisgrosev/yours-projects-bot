# -*- coding: utf-8 -*-
import os
import uuid
import sys
import time
import logging
import shutil
import re
import asyncio
from datetime import datetime
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from telegram import Bot, InputFile
import pdfplumber

import argparse

# === чат-история для DeepSeek ===
MAX_HISTORY = 10
chat_history = [
    {"role": "system", "content": "Отвечай только на русском языке. Четко следуй всем инструкциям."}
]

PROJECTS_DIR = "/app/data/files212/generate_project/projects"
os.makedirs(PROJECTS_DIR, exist_ok=True)
LOG_DIR = "/app/data/files212/generate_project/log"
os.makedirs(LOG_DIR, exist_ok=True)

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

parser = argparse.ArgumentParser(description="Генерация проекта в отдельном процессе и сопровождение пользователя в Telegram")
parser.add_argument('--token', required=True, help='Токен Telegram-бота')
parser.add_argument('--user_id', required=True, type=int, help='ID пользователя Telegram')
parser.add_argument('--fio_student', required=True, help='ФИО обучающегося')
parser.add_argument('--topic', required=True, help='Тема проекта')
parser.add_argument('--subject', required=True, help='Предмет')
parser.add_argument('--group', required=True, help='Группа')
parser.add_argument('--fio_teacher', required=True, help='ФИО преподавателя')
parser.add_argument('--num_points', required=True, type=int, help='Количество пунктов содержания (без учета источников)')
parser.add_argument('--spec_number', default='', help='Номер специальности')
parser.add_argument('--spec_name', default='', help='Название специальности')
parser.add_argument('--primer_path', required=True, help='Путь к шаблону primer.docx')
parser.add_argument('--output_dir', required=True, help='Папка для сохранения документов')
parser.add_argument('--deepseek_api_key', required=True, help='API ключ DeepSeek')
parser.add_argument('--admin_id', required=True, type=int, help='ID админа для отправки ошибок')

args = parser.parse_args()

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
log_filename = os.path.join(LOG_DIR, f"generate_project_log_{args.user_id}_{timestamp}.txt")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s %(message)s",
    handlers=[
        logging.FileHandler(log_filename, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

print("Process started with PID:", os.getpid())

# =================== КОНСТАНТЫ ==================
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
API_KEY = args.deepseek_api_key
HEADERS = {
    "Authorization": f"Bearer {API_KEY}",
    "Content-Type": "application/json"
}
ADMIN_ID = args.admin_id
USER_ID = args.user_id

def safe_soffice_convert(docx_path, pdf_path):
    import subprocess
    outdir = os.path.dirname(os.path.abspath(pdf_path))
    while True:
        try:
            print(f"Попытка конвертации: {docx_path} -> {pdf_path}")
            result = subprocess.run([
                '/usr/bin/soffice', '--headless', '--convert-to', 'pdf',
                '--outdir', outdir, docx_path
            ], check=True)
            if os.path.exists(pdf_path):
                print("Конвертация прошла успешно.")
                break
            else:
                print(f"PDF не найден: {pdf_path}")
                raise RuntimeError("PDF не найден после конвертации")
        except subprocess.CalledProcessError as e:
            print(f"Ошибка конвертации: {e}, код возврата: {e.returncode}")
            print("Жду 30 секунд и пробую снова...")
            time.sleep(30)
        except Exception as ex:
            print(f"Неожиданная ошибка: {ex}")
            time.sleep(30)

def make_doc_filename(
    fio: str,
    topic: str,
    user_id: int,
    timestamp: str,
    max_bytes: int = 180
) -> str:
    safe_fio = sanitize_filename(fio)
    safe_topic = sanitize_filename(topic)

    pretty_name = f"{safe_fio}. {safe_topic}.{timestamp}.docx"

    # считаем БАЙТЫ, а не символы
    if len(pretty_name.encode("utf-8")) > max_bytes:
        return f"project_{user_id}_{timestamp}.docx"

    return pretty_name

def sanitize_filename(text):
    forbidden_chars = '/\\:*?"<>|'
    for char in forbidden_chars:
        text = text.replace(char, '_')
    return text.strip()

def strip_leading_number(text):
    return re.sub(r"^\d+\.\s*", "", text)

def is_bold(fontname):
    return any(word in fontname.lower() for word in ["bold", "bd", "black", "heavy", "semibold"])

def fix_fonts(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(14)
                        run.font.name = "Times New Roman"

def remove_asterisks(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace("*", "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace("*", "")

def build_replacements(user_data):
    group = user_data['group']
    course = group[0] if group and group[0].isdigit() else ""
    spec_number = user_data.get('spec_number')
    spec_name = user_data.get('spec_name')
    if not spec_number or not spec_name:
        spec_number, spec_name = get_spec_by_group(group)
    return {
        "<<FIO>>": user_data['fio_student'],
        "<<THEME>>": user_data['topic'],
        "<<SUBJECT>>": user_data['subject'],
        "<<GROUP>>": group,
        "<<TEACHER>>": user_data['fio_teacher'],
        "<<COURSE>>": course,
        "<<SPECNUM>>": spec_number,
        "<<SPECTEXT>>": spec_name,
    }

def get_spec_by_group(group):
    group = group.upper()
    if "ТОД" in group:
        return "23.02.07", "Техническое обслуживание и ремонт двигателей, систем и агрегатов автомобилей"
    elif "ЭТ" in group:
        return "23.02.05", "Эксплуатация транспортного электрооборудования и автоматики (по видам транспорта, за исключением воздушного транспорта)"
    elif "СД" in group:
        return "08.02.12", "Строительство и эксплуатация автомобильных дорог, аэродромов и городских путей сообщения"
    elif "ОП" in group:
        return "23.02.01", "Организация перевозок и управление на транспорте (по видам)"
    else:
        return "", ""

def replace_placeholders_in_docx(doc, replacements):
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

def extract_clean_content(raw_text):
    logger.info("Обрабатываем текст...")
    list_items = re.findall(r'(?:\d+\.?|\-|\•)\s*.+', raw_text)
    if list_items:
        logger.info("Обнаружен список")
        return '\n'.join([item.strip() for item in list_items])
    return raw_text.strip()

def add_contents_page(doc, points):
    p_title = doc.add_paragraph()
    run = p_title.add_run("Содержание")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_title.paragraph_format.line_spacing = 1.5
    for idx, point in enumerate(points, 1):
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(18.5), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS
        )
        run = p.add_run(f"{idx}. {strip_leading_number(point).strip().rstrip('.')}\t")
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.line_spacing = 1.5

def add_page_numbers(doc, points):
    pid = os.getpid()
    unique_id = f"{pid}_{uuid.uuid4().hex}"
    temp_docx_path = f"_temp_toc_{unique_id}.docx"
    temp_pdf_path = f"_temp_toc_{unique_id}.pdf"
    logger.info(f"Сохраняем docx для постраничного анализа: {temp_docx_path}")
    doc.save(temp_docx_path)
    logger.info(f"Конвертируем docx в pdf: {temp_pdf_path}")
    safe_soffice_convert(temp_docx_path, temp_pdf_path)
    pages_dict = {}
    logger.info(f"Открываем pdf для анализа: {temp_pdf_path}")
    with pdfplumber.open(temp_pdf_path) as pdf:
        for point in points:
            title = strip_leading_number(point).strip().rstrip('.')
            normalized_title = " ".join(title.split()).lower()
            pages_dict[point] = None
            for i, page in enumerate(pdf.pages):
                if i < 2:
                    continue
                words = page.extract_words(extra_attrs=["fontname"])
                page_text = " ".join(w["text"] for w in words if is_bold(w.get("fontname", ""))).lower()
                normalized_page_text = re.sub(r"\s+", " ", page_text)
                if normalized_title in normalized_page_text:
                    pages_dict[point] = i + 1
                    break
            if pages_dict[point] is None:
                logger.warning(f"⚠️ Не найден заголовок: '{title}'")
    logger.info(f"Вставляем номера страниц в docx")
    for paragraph in doc.paragraphs:
        for idx, point in enumerate(points, 1):
            clean_label = f"{idx}. {strip_leading_number(point).strip().rstrip('.')}"
            if paragraph.text.startswith(clean_label):
                page = pages_dict.get(point)
                if page:
                    parts = paragraph.text.split('\t')
                    left = parts[0]
                    paragraph.clear()
                    run = paragraph.add_run(f"{left}\t{page}")
                    run.font.size = Pt(14)
                    run.font.name = "Times New Roman"
    for path in (temp_docx_path, temp_pdf_path):
        if os.path.exists(path):
            logger.info(f"Удаляем временный файл: {path}")
            os.remove(path)

def insert_page_break(paragraph):
    run = paragraph.insert_paragraph_before().add_run()
    run.add_break(WD_BREAK.PAGE)

def insert_page_break_after(paragraph):
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

def insert_page_break_after_last_content(doc, points):
    last_point_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        expected = f"{len(points)}."
        if paragraph.text.strip().startswith(expected):
            last_point_idx = idx
    if last_point_idx is not None:
        insert_page_break_after(doc.paragraphs[last_point_idx])

def add_page_breaks_around_contents(doc, points):
    contents_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Содержание":
            contents_idx = idx
            break
    if contents_idx is not None:
        insert_page_break(doc.paragraphs[contents_idx])
    last_point_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        for i, point in enumerate(points, 1):
            expected = f"{i}. {strip_leading_number(point).strip().rstrip('.')}"
            if paragraph.text.strip().startswith(expected):
                last_point_idx = idx
    if last_point_idx is not None:
        insert_page_break_after(doc.paragraphs[last_point_idx])

def send_deepseek_request(prompt, temperature=0.7, max_tokens=7000):
    global chat_history
    import requests

    # Добавить новое сообщение пользователя
    chat_history.append({"role": "user", "content": prompt})
    # Обрезать историю (system + последние MAX_HISTORY*2 сообщений)
    if len(chat_history) > 1 + MAX_HISTORY * 2:
        chat_history = [chat_history[0]] + chat_history[-MAX_HISTORY*2:]

    data = {
        "model": "deepseek-chat",
        "messages": chat_history,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    logger.info(f"Отправляем запрос в DeepSeek (чат): prompt={prompt[:50]}..., history_len={len(chat_history)}")
    response = requests.post(DEEPSEEK_API_URL, json=data, headers=HEADERS)
    response.raise_for_status()
    answer = response.json()['choices'][0]['message']['content']
    # Добавить ответ ассистента
    chat_history.append({"role": "assistant", "content": answer})
    logger.info(f"Ответ получен от DeepSeek (первые 50 символов): {answer[:50]}")
    return answer

async def send_deepseek_request_with_retry(prompt, temperature=0.7, max_tokens=7000, retries=3, delay=5):
    for attempt in range(1, retries + 1):
        try:
            logger.info(f"Попытка {attempt}: отправка запроса в DeepSeek")
            response = send_deepseek_request(prompt, temperature, max_tokens)
            return response
        except Exception as e:
            error_text = str(e)
            logger.error(f"Попытка {attempt}: Ошибка API: {error_text}")
            if "Response ended prematurely" in error_text or "Connection" in error_text or "timed out" in error_text:
                if attempt < retries:
                    logger.info(f"Повтор через {delay} секунд...")
                    await asyncio.sleep(delay)
                else:
                    logger.error("Все попытки исчерпаны, запрос не удался.")
                    raise
            else:
                raise

async def safe_send_message(bot, chat_id, *args, **kwargs):
    from telegram.error import TimedOut, NetworkError, TelegramError
    for i in range(20):
        try:
            logger.info(f"Отправляем сообщение [{i+1} попытка] пользователю {chat_id}: {args if args else ''} {kwargs if args else ''}")
            return await bot.send_message(chat_id, *args, **kwargs)
        except (TimedOut, NetworkError, TelegramError) as e:
            logger.warning(f"safe_send_message попытка {i+1}: {e}")
            await asyncio.sleep(2 * (i + 1))
    logger.error(f"Не удалось отправить сообщение после 20 попыток: {args}, {kwargs}")

async def main():
    logger.info(f"Запуск генератора с параметрами: {sys.argv}")
    bot = Bot(token=args.token)
    user_id = args.user_id
    user_data = {
        "fio_student": args.fio_student,
        "topic": args.topic,
        "subject": args.subject,
        "group": args.group,
        "fio_teacher": args.fio_teacher,
        "num_points": args.num_points,
        "spec_number": args.spec_number,
        "spec_name": args.spec_name
    }

    try:
        await safe_send_message(bot, user_id, "Генерируем пункты содержания...")
        topic = user_data['topic']
        subject = user_data['subject']
        num_points = user_data['num_points']

        content_prompt = (
            f"""Привет, я пишу проект по теме: {topic}, по предмету: {subject}.
            Составь нумерованный список из {num_points} уникальных, содержательных пунктов для содержания этого проекта. 
            В них не должно быть много текста, чтобы они поместились в содержание, в идеале около трех слов.
            Не добавляй подпунктов, пояснений, заголовков или инструкций — только сами пункты списка.
            Первый пункт должен быть по теме проекта, а не повторять формулировку задания.
            Первый пукт должен быть, "Введение", а последний "Заключение".
            Каждый пункт должен отражать отдельный аспект или раздел по теме.
            Оформи исключительно в виде нумерованного списка, без лишнего текста до и после."""
        )
        raw_content = await send_deepseek_request_with_retry(content_prompt)
        await safe_send_message(bot, user_id, "Обрабатываем текст...")

        clean_content = extract_clean_content(raw_content)
        await safe_send_message(bot, user_id, "Обнаружен список")

        points = clean_content.split("\n")[:num_points]  # только столько, сколько заказал пользователь
        points = [strip_leading_number(p).strip() for p in points if p.strip()]  # чистим

        # Безо всяких "Введение" и "Заключение" принудительно!
        points.append("Список используемых источников")  # только этот раздел принудительно

        texts = []
        MAX_RETRIES = 5

        for idx, point in enumerate(points, start=1):
            if point.lower() == "список используемых источников":
                await safe_send_message(bot, user_id, f"Генерируем {idx}/{len(points)}: Список источников...")

                sources_prompt = (
                    f"""Сформируй корректно оформленный по ГОСТу список из 7-10 используемых источников для проекта на тему: "{topic}" по предмету "{subject}".
                    Не добавляй заголовок, списки или пояснения — только сами записи источников, каждый с новой строки.
                    В ответе должны быть только сами записи источников, без лишнего текста до и после."""
                )
                for attempt in range(1, MAX_RETRIES + 1):
                    try:
                        sources_text = await send_deepseek_request_with_retry(sources_prompt)
                        texts.append(sources_text.strip())
                        break
                    except Exception as e:
                        if attempt == MAX_RETRIES:
                            texts.append("[Ошибка генерации источников. Попробуйте позже или обратитесь к поддержке. @denisgrosev]")
                        else:
                            await asyncio.sleep(5)
            else:
                await safe_send_message(bot, user_id, f"Генерируем текст для пункта {idx}/{len(points)}...")
                text_prompt = (
                    f"""Напиши развернутый текст объёмом примерно 420 слов на тему: "{point}".
                    Пиши цельный, связный и информативный текст, избегая повторов и "воды".
                    Не используй подзаголовки, маркированные или нумерованные списки, таблицы, цитаты и выделения.
                    Не начинай предложения с дефиса, тире, точки или других символов, не соответствующих обычному началу предложения.
                    Излагай информацию в логической последовательности, плавно переходя от одной мысли к другой.
                    Текст должен быть написан на русском языке и подходить для включения в основную часть научного или учебного проекта.
                    В ответе должен быть только сплошной текст, без каких-либо дополнительных инструкций, пояснений или рамок.
                    """
                )
                for attempt in range(1, MAX_RETRIES + 1):
                    try:
                        raw_text = await send_deepseek_request_with_retry(text_prompt)
                        texts.append(raw_text.strip())
                        break
                    except Exception as e:
                        if attempt == MAX_RETRIES:
                            texts.append("[Ошибка генерации текста. Попробуйте позже или обратитесь к поддержке. @denisgrosev]")
                        else:
                            await asyncio.sleep(5)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        primer_doc_path = args.primer_path
        output_dir = args.output_dir

        logger.info(f"Проверка наличия шаблона: {primer_doc_path}")
        if not os.path.isfile(primer_doc_path):
            logger.error(f"Файл шаблона не найден: {primer_doc_path}")
            raise FileNotFoundError(f"Файл шаблона не найден: {primer_doc_path}")

        logger.info(f"Создание папки для документов: {output_dir}")
        os.makedirs(output_dir, exist_ok=True)
        replacements = build_replacements(user_data)

        from io import BytesIO
        with open(primer_doc_path, "rb") as f:
            primer_bytes = f.read()
        mem_doc = BytesIO(primer_bytes)
        doc = Document(mem_doc)

        replace_placeholders_in_docx(doc, replacements)
        fix_fonts(doc)
        add_contents_page(doc, points)
        for idx, (point, text) in enumerate(zip(points, texts), 1):
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run(point)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing = 1.5

            p2 = doc.add_paragraph(text)
            if point.lower() == "список используемых источников":
                # Без отступа, строго по левому краю!
                p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p2.paragraph_format.line_spacing = 1.5
                p2.paragraph_format.first_line_indent = Cm(0)
            else:
                p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p2.paragraph_format.line_spacing = 1.5
                p2.paragraph_format.first_line_indent = Cm(1.27)
            for run in p2.runs:
                run.font.size = Pt(14)
                run.font.name = "Times New Roman"
        remove_asterisks(doc)
        add_page_breaks_around_contents(doc, points)
        add_page_numbers(doc, points)
        insert_page_break_after_last_content(doc, points)

        output_buffer = BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        
        filename = make_doc_filename(
            fio=user_data["fio_student"],
            topic=user_data["topic"],
            user_id=user_id,
            timestamp=timestamp
        )
        
        doc_filename = os.path.join(output_dir, filename)
        logger.info(f"Сохраняем финальный docx на диск: {doc_filename}")
        doc.save(doc_filename)
        
        project_copy_path = os.path.join(PROJECTS_DIR, os.path.basename(doc_filename))
        shutil.copyfile(doc_filename, project_copy_path)
        logger.info(f"Проект продублирован в: {project_copy_path}")
        
        await safe_send_message(bot, user_id, "Проект успешно создан! Документ отправлен в чат.")
        
        await bot.send_document(
            user_id,
            InputFile(output_buffer, filename=filename),
            caption="Спасибо за покупку :З"
        )
    except Exception as e:
        logger.error(f"Exception в генераторе: {e}", exc_info=True)
        try:
            await safe_send_message(bot, ADMIN_ID, f"Ошибка генерации у пользователя {user_id}: {e}")
            await safe_send_message(bot, user_id, "Произошла ошибка при генерации проекта. Сообщите @denisgrosev.")
        except Exception as err:
            logger.error(f"Ошибка при отправке ошибки админу: {err}", exc_info=True)

if __name__ == "__main__":
    asyncio.run(main())

